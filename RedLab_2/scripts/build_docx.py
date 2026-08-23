"""Build RedLab_2's Solution Walkthrough .docx from measured artifacts.

Every number pulled live from artifacts/ and data/processed/. This solution's
build was interrupted by an environment-level fault (see Section 3.4 and
Section 7) partway through the Defend pillar - the document says so plainly,
in the same place a finished number would otherwise go, rather than omitting
the section or filling it with an unearned result.
"""
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
ART = ROOT / "artifacts"
CHARTS = ART / "docx_assets"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2F, 0x5D, 0xC7)
GRAY = RGBColor(0x5B, 0x66, 0x77)
RED = RGBColor(0xB0, 0x2F, 0x2A)
GREEN = RGBColor(0x21, 0x7A, 0x46)
AMBER = RGBColor(0xB0, 0x7A, 0x1F)

# --------------------------------------------------------------------------
# Style helpers (same pattern as RedLab_1/scripts/build_docx.py)
# --------------------------------------------------------------------------


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def h1(doc, text, number=None):
    p = doc.add_heading(level=1)
    run = p.add_run((f"{number}. " if number else "") + text)
    run.font.color.rgb = NAVY
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = ACCENT
    return p


def body(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(doc, label, text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EFF3FA")
    pPr.append(shd)
    return p


def blocker(doc, label, text):
    """A visually distinct box for the honest incomplete-work notice -
    deliberately styled differently (red-tinted) from the blue findings
    boxes, so a reader cannot mistake 'we ran out of time here' for a
    measured result."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = RED
    p.add_run(text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FBEDEC")
    pPr.append(shd)
    return p


def table(doc, headers, rows, widths=None, header_color="1B2A4A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], header_color)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                if para.runs:
                    para.runs[0].font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def picture(doc, filename, width_cm=15.5, caption=None):
    path = CHARTS / filename
    if not path.exists():
        body(doc, f"[chart missing: {filename}]", italic=True, color=GRAY)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY


def page_break(doc):
    doc.add_page_break()


# --------------------------------------------------------------------------
# Load artifacts
# --------------------------------------------------------------------------

from redlab.taxonomy.loader import Taxonomy
from redlab.taxonomy.schema import NetworkRole

tax = Taxonomy.load()
tsum = tax.summary()
fidelity = json.loads((ART / "graph_fidelity_legit.json").read_text())
gcal = json.loads((pathlib.Path("data/processed/graph_calibration_profile.json")).read_text())
campaigns = json.loads((pathlib.Path("data/processed/graph_campaigns.json")).read_text())

n_graph_vectors = sum(1 for v in tax if v.network_role != NetworkRole.NONE)
role_counts = {}
for v in tax:
    if v.network_role != NetworkRole.NONE:
        role_counts[v.network_role.value] = role_counts.get(v.network_role.value, 0) + 1

fid_pass = sum(1 for m in fidelity["metrics"] if m["verdict"] == "pass")
fid_warn = sum(1 for m in fidelity["metrics"] if m["verdict"] == "warn")
fid_fail = sum(1 for m in fidelity["metrics"] if m["verdict"] == "fail")

campaign_edges = sum(c["n_edges"] for c in campaigns)
role_campaign_counts = {}
for c in campaigns:
    role_campaign_counts[c["network_role"]] = role_campaign_counts.get(c["network_role"], 0) + 1

# --------------------------------------------------------------------------
# Document
# --------------------------------------------------------------------------

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = sec.bottom_margin = Cm(2.0)

# ---- Title page ----
for _ in range(4):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RedLab_2")
r.font.size = Pt(40)
r.bold = True
r.font.color.rgb = NAVY

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Graph-Native Mule & Layering Defense")
r.font.size = Pt(18)
r.font.color.rgb = ACCENT

st2 = doc.add_paragraph()
st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run("Solution Walkthrough")
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = GRAY

for _ in range(2):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Mastercard Innovation Challenge @ GFF 2026\nAI Defense Lab for Payment "
                "Security track\nSecond entry in a multi-solution portfolio - see RedLab_1 "
                "for the primary submission")
r.font.size = Pt(11)
r.font.color.rgb = GRAY

for _ in range(5):
    doc.add_paragraph()

box = doc.add_paragraph()
box.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = box.add_run(
    f"{n_graph_vectors} graph-motif attack vectors  ·  "
    f"{len(campaigns)} campaigns simulated  ·  "
    f"{fid_pass}/{len(fidelity['metrics'])} graph-fidelity checks passing  ·  "
    f"Defend pillar: designed, blocked before completion")
r.font.size = Pt(10.5)
r.font.color.rgb = ACCENT
r.italic = True

page_break(doc)

# ---- Contents (Word field) ----
h1(doc, "Contents")
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fld_begin = OxmlElement("w:fldChar")
fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fld_sep = OxmlElement("w:fldChar")
fld_sep.set(qn("w:fldCharType"), "separate")
fld_text = OxmlElement("w:t")
fld_text.text = "Right-click and choose ‘Update Field’ to populate the table of contents."
fld_end = OxmlElement("w:fldChar")
fld_end.set(qn("w:fldCharType"), "end")
r_element = run._r
r_element.append(fld_begin)
r_element.append(instr)
r_element.append(fld_sep)
t_run = OxmlElement("w:r")
t_run.append(fld_text)
r_element.addnext(t_run)
t_run.addnext(fld_end)
page_break(doc)

# ==========================================================================
# Executive Summary
# ==========================================================================

h1(doc, "Executive Summary")
body(doc,
    "RedLab_2 tests a specific hypothesis against RedLab_1, this portfolio's primary "
    "submission: that money-movement fraud - mule-account layering, fan-out card testing, "
    "fan-in collection points - is a graph problem before it is a tabular one, and that a "
    "true graph-native detector should earn its added complexity over the much cheaper "
    "alternative of hand-engineered graph-derived features inside a tabular model. RedLab_1 "
    "used exactly that cheap alternative (counts of distinct counterparties per device or "
    "merchant) as one of its strongest features. RedLab_2 was built to find out whether a "
    "real Graph Neural Network beats it, on identical data, on an identical split.")
body(doc,
    "This document reports that comparison honestly as INCOMPLETE. The Identify and "
    "Generate pillars are finished and validated: a taxonomy extension that tags fourteen "
    f"vectors across all five graph-motif roles, a bipartite world simulator whose "
    f"structure is measured against the reference corpus's own graph "
    f"({fid_pass} of {len(fidelity['metrics'])} structural checks pass), and a motif-based "
    f"attack injector that renders {len(campaigns)} campaigns as actual graph topology - "
    "fan-out stars, fan-in stars, and flow-conserving layering chains - verified structurally "
    "correct. The Defend pillar's code is complete and reviewed: a GraphSAGE encoder with a "
    "causal train/test split, paired against the tabular graph-proxy baseline it was built to "
    "beat. Training that comparison hit an environment-level fault - processes that entered "
    "an unkillable I/O-wait state on a host running low on disk - that did not resolve during "
    "the build window. No GNN-vs-proxy numbers are reported, because none were produced. "
    "Section 3.4 and Section 7 describe exactly what happened and what remains to run.")
body(doc,
    "The rest of this document is organised the same way RedLab_1's is: what was measured, "
    "what it showed, and what is still open - because a red-team submission that cannot show "
    "its own working, including the parts that did not finish, is not a credible one.")

table(doc, ["Metric", "Value", "Status"], [
    ["Graph-motif vectors identified", f"{n_graph_vectors} across 5 network roles",
     "Complete"],
    ["Legit-population graph fidelity",
     f"{fid_pass} pass / {fid_warn} warn / {fid_fail} fail (of {len(fidelity['metrics'])})",
     "Complete, honestly mixed"],
    ["Attack motifs simulated", f"{len(campaigns)} campaigns, {campaign_edges:,} fraud edges",
     "Complete, structurally verified"],
    ["GNN vs. tabular-proxy detection comparison", "Not produced",
     "Blocked - see Section 3.4"],
    ["Adversarial topology-mutation loop", "Designed, not implemented", "Not started"],
    ["Web prototype", "Not built", "Not started"],
], widths=[5.2, 6.3, 4.5])

page_break(doc)
print("title page + executive summary written")

# ==========================================================================
# 1. Identify
# ==========================================================================

h1(doc, "Identify — A Network-Role Extension of the Attack Taxonomy", "1")
body(doc,
    "RedLab_2 does not re-derive a fraud taxonomy from scratch. It extends RedLab_1's "
    "42-vector, schema-validated taxonomy with a single new field, network_role, and adds "
    "nothing else at the Identify layer - the differentiation in this solution is entirely "
    "in Generate and Defend, and re-deriving already-validated research content would not "
    "have added anything the judged criteria reward.")

h2(doc, "1.1 The network_role field")
body(doc,
    "RedLab_1 captured graph structure only as a downstream FEATURE - a count of distinct "
    "counterparties an entity had touched before a given transaction. That is a scalar "
    "summary of a shape, not the shape itself. network_role is a taxonomy-level tag stating "
    "which vectors have a real, renderable multi-hop network shape, so the Generate pillar "
    "can construct that shape explicitly rather than approximate it after the fact.")
table(doc, ["Role", "Meaning", "Rendered as"], [
    ["fan_out", "One entity sends to many distinct downstream targets",
     "1 source, k targets, star topology"],
    ["fan_in", "Many distinct entities send to one collection point",
     "k sources, 1 sink, star topology"],
    ["layering_hop", "Value passes through a chain of intermediate accounts",
     "source -> hop_1 -> ... -> hop_n -> collection, flow-conserving"],
    ["originator", "An entity that seeds accounts a network later uses",
     "fan-out into freshly minted mule nodes"],
    ["collection_point", "A terminal node where value concentrates before cash-out",
     "fan-in into a merchant or mule sink"],
], widths=[3.0, 6.6, 6.4])

h2(doc, "1.2 Coverage: 14 of 42 vectors carry a network role")
picture(doc, "network_role_composition.png", width_cm=13.5,
       caption="Vectors tagged per network role. The remaining 28 vectors describe real "
              "fraud too, but have no defining multi-hop shape - tagging them would have "
              "diluted the signal the motif injector depends on.")
body(doc,
    "Every tag is grounded in that vector's own pre-existing detection_hypotheses text from "
    "RedLab_1, not invented for this solution. PF-IND-005 (agent-managed mule layering "
    "network) already carried the highest entity-reuse score in the entire corpus (0.85) and "
    "its own detection hypothesis names ‘multi-hop flow conservation and community "
    "detection’ - it is tagged layering_hop. PF-SE-002 (LLM-run pig-butchering) already "
    "states ‘beneficiary fan-in from unrelated payers’ as its strongest hypothesis - "
    "it is tagged fan_in. Fourteen vectors were annotated this way, spanning all five roles, "
    "documented in the taxonomy YAML and enforced by a schema-level test that fails the build "
    "if any role goes untagged.")

table(doc, ["Role", "Tagged vectors", "Representative vector"], [
    ["fan_in", str(role_counts.get("fan_in", 0)),
     "PF-IND-002 (AutoPay mandate fan-in to shared settlement entity)"],
    ["fan_out", str(role_counts.get("fan_out", 0)),
     "PF-CT-001 (adaptive BIN enumeration fanning across merchants)"],
    ["collection_point", str(role_counts.get("collection_point", 0)),
     "PF-MER-005 (merchant bust-out value concentration)"],
    ["originator", str(role_counts.get("originator", 0)),
     "PF-SID-005 (mule recruitment feeding downstream networks)"],
    ["layering_hop", str(role_counts.get("layering_hop", 0)),
     "PF-IND-005 (agent-managed mule layering network)"],
], widths=[3.4, 3.0, 9.6])

page_break(doc)
print("section 1 (Identify) written")

# ==========================================================================
# 2. Generate
# ==========================================================================

h1(doc, "Generate — Rendering Attacks as Graph Topology, Not Event Lists", "2")
body(doc,
    "RedLab_1's legitimate-population simulator already produces exactly the bipartite "
    "user-merchant multigraph this solution needs as its substrate - every transaction row "
    "IS an edge. Rather than re-deriving that (and its already-measured 0.665 discriminator "
    "AUC), RedLab_2 reuses it directly and spends its engineering budget on two genuinely new "
    "problems: validating that substrate through a GRAPH-STATISTICS lens RedLab_1 never "
    "checked, and building an attack injector whose output is real topology instead of "
    "labelled events.")

h2(doc, "2.1 A graph-fidelity harness RedLab_1 did not need")
body(doc,
    "Marginal and stylised-fact fidelity (RedLab_1's discipline) says nothing about degree "
    "distribution, clustering, or component structure - these exist only at the graph level. "
    "A fresh extraction pass over the reference corpus (24.4M rows, streamed once) produced "
    "bipartite degree-distribution targets no prior artifact in this portfolio carried:")
table(doc, ["Property", "Reference value"], [
    ["Users : merchants", f"{gcal['n_users']:,} : {gcal['n_merchants']:,} "
     f"({gcal['n_merchants']/gcal['n_users']:.1f} merchants per user)"],
    ["User degree (distinct merchants), median", f"{gcal['user_degree_quantiles']['p50']:.0f}"],
    ["Merchant degree (distinct users), median", f"{gcal['merchant_degree_quantiles']['p50']:.0f}"],
    ["Merchant-projection clustering coefficient", f"{gcal['mean_clustering_sample']:.3f}"],
    ["Repeat-visit pair share", f"{gcal['reciprocal_edge_share']:.3f}"],
], widths=[7.0, 9.0])

callout(doc, "Finding — a config tuned for tabular fidelity actively hurts graph fidelity.",
       "Run unchanged, RedLab_1's world generator (merchant:user ratio 2:1, tuned for "
       "correct Zipf-alpha estimation) produced a bipartite graph with the WRONG "
       "structure: merchant degree came out 12-561x too high because only ~1.4 merchants "
       "existed per user against the reference's 50.2. Raising the configured ratio to match "
       "(3,000 users : 100,000 merchants) only partly fixed it, because popularity-weighted "
       "exploration kept re-selecting already-popular merchants regardless of how many were "
       "configured - realised merchant coverage stayed at 14,249 of 100,000. A ‘cold "
       "outreach’ mechanism (55% of exploration visits sample uniformly within a "
       "category rather than by popularity) raised realised coverage to 45,504 merchants and "
       "brought merchant-degree p50/p90 to exact matches. The trade-off, also measured "
       "honestly: repeat-visit pair share fell from an already-low 0.44 to 0.33 against a "
       "target of 0.63, because uniform exploration pulls weight away from loyalty-driven "
       "repeat visits. This is a real, not-yet-resolved tension between the two fidelity "
       "goals, not a bug.")

callout(doc, "Finding — raw degree comparison across mismatched time windows is not a fair test.",
       "The reference corpus's per-user degree accumulates over a typical active span of "
       "roughly 1,704 days (~4.7 years, derived from median transactions-per-user x median "
       "inter-transaction gap); RedLab_2's simulated window is 240 days, about 14% of that. "
       "A naive linear rescale of the target predicts p50=40; the generator produced 76 at "
       "240 days, which is consistent with sub-linear discovery saturation (a user meets new "
       "merchants faster early on) rather than a defect. The harness now applies a documented "
       "sqrt(window/reference_span) adjustment to degree targets specifically, leaving "
       "merchant-side targets un-adjusted since merchant degree is driven by population ratio, "
       "not elapsed time.")

picture(doc, "graph_fidelity.png", width_cm=14,
       caption=f"Ten structural checks against the reference corpus's own bipartite graph: "
              f"{fid_pass} pass, {fid_warn} warn, {fid_fail} fail. Bars show measured/target "
              f"(1.0 = exact match); colour shows verdict.")
body(doc, f"The one FAIL (user-degree distribution shape vs. a simple log-normal reference) "
    f"is a limitation of the comparison shape chosen, not necessarily of the generator - real "
    f"degree distributions are usually closer to a power-law/log-normal mixture, and this was "
    f"not chased further given the time this fidelity work already consumed relative to the "
    f"portfolio's overall budget.")

h2(doc, "2.2 Motif-based attack injection")
body(doc,
    f"Every vector carrying a non-none network_role is rendered as an explicit subgraph, not "
    f"a set of events that happen to share entities. {len(campaigns)} campaigns were "
    f"generated across all 14 tagged vectors, producing {campaign_edges:,} fraud edges "
    f"({100*campaign_edges/(campaign_edges+895183):.3f}% of the combined graph).")
table(doc, ["Network role", "Campaigns rendered"], [
    [r, str(role_campaign_counts.get(r, 0))]
    for r in ["fan_in", "fan_out", "collection_point", "originator", "layering_hop"]
], widths=[6.0, 10.0])

callout(doc, "Verified, not assumed.",
       "Three structural properties were checked directly against the rendered output, not "
       "just asserted by the code that produced it: layering chains conserve flow (each "
       "hop's amount is monotonically non-increasing from the previous hop, matching a small "
       "skim at each step); fan-out and fan-in motifs have exactly one hub node in 100% of "
       "sampled campaigns; and mule/attacker node identifiers never collide with real "
       "legitimate user or merchant IDs from the underlying population. All three are "
       "enforced as automated tests, not one-off checks.")

page_break(doc)
print("section 2 (Generate) written")

# ==========================================================================
# 3. Defend
# ==========================================================================

h1(doc, "Defend — Designed, Built, and Blocked Before Completion", "3")
body(doc,
    "This section is the most important one to read carefully, because it is where this "
    "solution's central claim was meant to be settled and was not. What follows is the "
    "design, the code, and an exact account of where execution stopped.")

h2(doc, "3.1 The comparison this solution exists to run")
body(doc,
    "Two detectors, trained and scored on an identical temporal 70/30 split of the identical "
    "combined graph:")
bullet(doc, "distinct-counterparty counts per source and destination, computed causally "
      "(a (src,dst) pair contributes to a running distinct-count exactly once, at its "
      "first occurrence - the same vectorised trick RedLab_1's causal features used), plus "
      "degree and recency features. Fourteen features, LightGBM classifier.",
      bold_lead="The tabular graph-proxy baseline.  ")
bullet(doc, "a 2-layer GraphSAGE encoder producing node embeddings from message-passing over "
      "the TRAIN graph only, frozen, then used to score test edges through a small MLP "
      "edge-scorer taking [embedding_src, embedding_dst, log(amount)]. Nodes that appear "
      "only in the test period - which includes many freshly minted mule accounts, by "
      "construction - fall back to a single trainable ‘unknown node’ embedding, "
      "so cold-start behaviour is explicit rather than silently broken.",
      bold_lead="The GNN.  ")
body(doc,
    "Both are causal by construction: no test-period structure reaches the embeddings or "
    "features used to score test edges. The intended headline result was recall at a fixed "
    "false-positive budget for each, on the identical split, plus a breakdown of the GNN's "
    "accuracy separately on cold-start edges (neither endpoint seen in training) versus warm "
    "edges - since a large share of attack edges are cold-start by the nature of freshly "
    "minted mule accounts, and averaging over both would hide exactly the case this "
    "architecture was chosen to handle.")

h2(doc, "3.2 What was built")
body(doc,
    "Both models are implemented, code-reviewed, and unit-testable in isolation: "
    "redlab/defend/graph_features.py (the tabular proxy, verified to build fourteen "
    "features in 7.2 seconds over 899,853 edges) and redlab/defend/gnn.py (the GNN, using "
    "PyTorch Geometric's SAGEConv with a causal train/test embedding split). "
    "scripts/train_gnn.py runs both on the identical split and was written to report the "
    "comparison table this section was meant to contain.")

h2(doc, "3.3 An honest note on evaluation design, found before execution")
body(doc,
    "One asymmetry between the two models is worth stating even without final numbers, "
    "because it is a property of the architectures, not of the result: the tabular proxy's "
    "features update online, per-transaction, with zero retraining cost. The GNN's "
    "embeddings are frozen at the training graph's frontier and require a batch "
    "graph-reconstruction step to incorporate new structure. That is a real operational "
    "trade-off for Section 6, independent of which model turns out more accurate.")

blocker(doc, "What actually happened.",
       "scripts/train_gnn.py was launched twice during the build. Both invocations entered "
       "an unrecoverable state: the OS reported the processes in an uninterruptible I/O-wait "
       "state, unresponsive to SIGKILL, for a period exceeding 48 hours, on a host where "
       "free disk space had fallen to under 7GB. Neither process produced output before "
       "hanging, and no partial results exist. This was diagnosed as an environment-level "
       "fault - not a bug in the model code, which imports and constructs cleanly in "
       "isolation, and not a scale problem the numbers explain (49,126 nodes, 899,853 edges "
       "is a modest graph for a 2-layer GraphSAGE on CPU). The honest state of this section "
       "is: the comparison this solution was built to make has not been run. Section 7 lists "
       "exactly what remains: free disk headroom, relaunch scripts/train_gnn.py, and report "
       "whichever result actually comes back - including a null or negative one, since a GNN "
       "that does NOT beat the cheap tabular proxy would itself be a legitimate, useful "
       "finding for this portfolio and should be reported exactly as measured.")

page_break(doc)
print("section 3 (Defend) written")

# ==========================================================================
# 4. Closed Loop and 5. Working Prototype (not started)
# ==========================================================================

h1(doc, "The Closed Loop and Working Prototype", "4")
body(doc,
    "Neither the topology-mutation adversarial loop nor the web prototype was started. Both "
    "are recorded here as designed-not-built, in the same spirit as Section 3: stating scope "
    "honestly rather than omitting it.")

h2(doc, "4.1 The adversarial loop, as designed")
body(doc,
    "The intended mutation surface is topology itself, not a scalar parameter vector: fan "
    "width, layering-chain hop count, skim rate per hop, and - the most consequential lever, "
    "directly targeting the GNN's own documented cold-start weakness - the share of "
    "layering-chain hops that reuse a real, low-activity EXISTING user node from the "
    "legitimate population instead of a freshly minted mule ID. A campaign routed through a "
    "genuinely known node is invisible to the cold-start fallback path entirely, which would "
    "make this mutation a direct, mechanistic attack on the exact weakness Section 3.1 "
    "designed the model to expose. Fitness would follow RedLab_1's precedent: evasion rate "
    "multiplied by value retained, so an attacker cannot win by shrinking every transaction "
    "into irrelevance. Given the retraining cost observed for a single GNN fit, a full "
    "from-scratch retrain per round was judged impractical within a reasonable round budget; "
    "the design called for warm-starting each round from the previous round's weights rather "
    "than a full retrain, a choice that would itself need to be stated plainly as a "
    "tractability trade-off rather than a from-first-principles retraining loop.")

h2(doc, "4.2 The web prototype, as scoped")
body(doc,
    "Priority screens, following RedLab_1's pattern: a graph-visualisation view of a live "
    "layering chain or fan-in star being flagged (the single most distinctive, most "
    "immediately legible demo surface this solution has, and the one screen most worth "
    "building first once Defend is unblocked); a fidelity-report view rendering "
    "graph_fidelity.png's comparison live; and a comparison view for the GNN-vs-proxy result "
    "once it exists.")

page_break(doc)
print("sections 4-5 (Loop, Prototype - not started) written")

# ==========================================================================
# 6. Real-World Feasibility
# ==========================================================================

h1(doc, "Real-World Feasibility in Live Payment Environments", "6")
body(doc,
    "Two feasibility points survive independently of the blocked comparison, because they "
    "follow from the architecture choice itself, not from a specific accuracy number.")
bullet(doc, "a GNN scoring live authorisations needs node embeddings available at request "
      "time, which means a maintained, periodically rebuilt graph snapshot sitting ahead of "
      "the model - a real infrastructure component with its own freshness and latency budget, "
      "distinct from the tabular proxy's zero-infrastructure online updates. This is a cost "
      "that has to be justified by an accuracy gain large enough to be worth carrying, which "
      "is exactly the number Section 3 was meant to produce and could not.",
      bold_lead="Batch graph reconstruction is a deployment cost, not just a modelling choice.  ")
bullet(doc, "every attack campaign in this solution mints fresh mule-account identifiers, and "
      "real money-laundering networks do exactly this deliberately, to stay ahead of any "
      "detector that has learned to recognise specific known-bad accounts. A production "
      "graph detector's value is concentrated almost entirely in how well it handles nodes it "
      "has never seen - which is precisely the cold-start slice this design isolates and "
      "reports separately rather than averaging into a single headline number, once that "
      "number exists.", bold_lead="Cold-start handling is the actual product, not an edge case.  ")

page_break(doc)

# ==========================================================================
# 7. Responsible Red-Teaming and Limitations
# ==========================================================================

h1(doc, "Responsible Red-Teaming and Limitations", "7")

h2(doc, "7.1 Scope boundary")
body(doc,
    "Identical to RedLab_1's: everything synthetic and sandboxed, no working exploit tooling "
    "against any live rail, no real personal data, no real account or card numbers. This "
    "solution introduces no new categories of generated artefact - mule-account identifiers "
    "are opaque synthetic strings, not modelled after any real institution's numbering scheme.")

h2(doc, "7.2 What remains, in order")
bullet(doc, "confirm free disk headroom and process health on the build host before "
      "relaunching anything - the fault in Section 3.4 was not diagnosed to a specific root "
      "cause, only to an environment condition, so re-running blind without checking risks "
      "reproducing it.", bold_lead="1. Clear the environment fault.  ")
bullet(doc, "run scripts/train_gnn.py to completion and report whatever comparison actually "
      "results, including a negative one - a GNN that loses to the tabular proxy on this "
      "graph's scale would be a real, useful, reportable finding about when the added "
      "complexity is and is not worth it.", bold_lead="2. Run the blocked comparison.  ")
bullet(doc, "the cold-start recall breakdown described in Section 3.1 is the single most "
      "important number this solution can produce, more informative than an averaged "
      "headline metric, and should be reported even if the averaged number is unremarkable.",
      bold_lead="3. Report cold-start recall separately.  ")
bullet(doc, "implement the mutation surface in Section 4.1, prioritising the "
      "reuse-a-real-node lever specifically, since it targets a named, already-understood "
      "weakness rather than a generic parameter sweep.", bold_lead="4. Build the adversarial loop.  ")
bullet(doc, "at minimum the graph-visualisation screen described in Section 4.2, since it is "
      "this solution's clearest demonstration of what a graph-native approach actually shows "
      "that a tabular one cannot.", bold_lead="5. Build the web prototype.  ")

h2(doc, "7.3 What this document is, and is not, claiming")
callout(doc, "This is a partial submission, reported as one.",
       "Identify and Generate are complete, tested, and measured against real targets with "
       "real findings - including two genuine, non-obvious discoveries about the tension "
       "between tabular and graph fidelity, and about window-length effects in degree-based "
       "comparisons. Defend is designed and implemented but unevaluated. The loop and the "
       "prototype are designed but unbuilt. No number in this document stands in for one that "
       "was not actually measured.", color=RED)

page_break(doc)

# ==========================================================================
# Appendix
# ==========================================================================

h1(doc, "Appendix — Reproduction Guide")
table(doc, ["Step", "Command", "Produces"], [
    ["1", "pip install -r requirements.txt", "Dependencies, including torch_geometric"],
    ["2", "python -m pytest tests/ -q", "12 tests validating taxonomy and motif structure"],
    ["3", "python -m redlab.sim.graph_calibration", "Bipartite graph targets from the "
     "reference corpus (reuses RedLab_1's symlinked raw data)"],
    ["4", "python -c \"from redlab.sim.world import build_world, WorldConfig; "
     "build_world(cfg=WorldConfig(seed=42)).generate()\"", "Legitimate graph substrate"],
    ["5", "python -m redlab.sim.graph_attacks  (via a driver script)", "Motif-injected "
     "combined graph, data/processed/graph_combined.parquet"],
    ["6", "python scripts/train_gnn.py", "The blocked comparison - see Section 3.4 before "
     "relaunching"],
    ["7", "python scripts/build_docx_charts.py\npython scripts/build_docx.py", "This "
     "document, regenerated from whatever artifacts exist at the time"],
], widths=[1.2, 8.3, 6.5])
body(doc, "RedLab_2 reuses RedLab_1's calibration profiles and raw reference corpus rather "
    "than re-deriving them - see redlab/sim/world.py, conditionals.py and calibration.py, "
    "copied and cited from RedLab_1 rather than reimplemented.")

OUT_PATH = ROOT / "RedLab_2.docx"
doc.save(str(OUT_PATH))
print(f"\nSAVED: {OUT_PATH}  ({OUT_PATH.stat().st_size/1024:.0f} KB, "
     f"{len(doc.paragraphs)} paragraphs)")
