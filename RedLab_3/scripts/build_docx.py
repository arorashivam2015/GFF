"""Build RedLab_3's short-form solution walkthrough. Deliberately compact,
matching this solution's scope."""
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2F, 0x5D, 0xC7)
GRAY = RGBColor(0x5B, 0x66, 0x77)
RED = RGBColor(0xB0, 0x2F, 0x2A)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def h1(doc, text, number=None):
    p = doc.add_heading(level=1)
    r = p.add_run((f"{number}. " if number else "") + text)
    r.font.color.rgb = NAVY
    return p


def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def callout(doc, label, text, color=ACCENT, fill="EFF3FA"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(text)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p._p.get_or_add_pPr().append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, "1B2A4A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            if cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


# --------------------------------------------------------------------------
from redlab.taxonomy.loader import Taxonomy

tax = Taxonomy.load()
swarm = json.loads((ROOT / "artifacts" / "swarm_summary.json").read_text())
campaigns = json.loads((ROOT / "data" / "processed" / "generated_campaigns.json").read_text())
generated_text = next(c for c in campaigns if c["vector_id"] in swarm["generated_vectors"])

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = sec.bottom_margin = Cm(2.0)

for _ in range(5):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RedLab_3")
r.font.size = Pt(38)
r.bold = True
r.font.color.rgb = NAVY
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("LLM Red-Team Swarm")
r.font.size = Pt(17)
r.font.color.rgb = ACCENT
st2 = doc.add_paragraph()
st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run("Solution Walkthrough (short form)")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = GRAY
for _ in range(2):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Mastercard Innovation Challenge @ GFF 2026\n"
                "One entry in a multi-solution portfolio - see RedLab_1 (primary), "
                "RedLab_2, RedLab_6 and RedLab_7 for the other entries")
r.font.size = Pt(10.5)
r.font.color.rgb = GRAY
doc.add_page_break()

# ==========================================================================
h1(doc, "Summary")
body(doc,
    "This solution's premise: the GenAI uplift in modern payment fraud lives in the LANGUAGE "
    "and the PLAN - the pretext, the negotiation, the persona - not just in the transaction "
    "numbers. The design called for an LLM Proposer agent authoring campaign scripts, "
    "conditioned on the taxonomy's own vector descriptions.")
body(doc,
    "What actually happened is the finding: of five generation attempts, spanning the "
    "taxonomy's social-engineering and agentic-commerce families, the model refused four. "
    "This was not worked around, rephrased, or retried to extract compliance - the refusals "
    "are reported as the primary result, because they are real, measured evidence about the "
    "practical feasibility of this entire solution category, which is exactly what the "
    "brief's real-world-feasibility criterion asks for.")

table(doc, ["Metric", "Value"], [
    ["Generation attempts (one call per vector, via the Claude CLI)", "5"],
    ["Refused", f"{swarm['n_refused']} of 5"],
    ["Generated", f"{swarm['n_generated']} of 5"],
    ["Method", "Non-interactive `claude -p`, no separate Critic call (see Limitations)"],
], widths=[11.0, 5.0])

h1(doc, "What's reused vs. what's new", "1")
bullet(doc, "taxonomy (42 vectors, unchanged), calibrated world simulator, attack injection "
      "- copied from RedLab_1 without modification, providing the numeric transaction side.",
      bold_lead="Reused.  ")
bullet(doc, "the Proposer (redlab/sim/llm_swarm.py) - one LLM call per selected vector, "
      "grounded in that vector's own summary and preconditions from the RedLab_1 taxonomy.",
      bold_lead="New.  ")

h1(doc, "The generation attempt", "2")
body(doc,
    "Five vectors were selected across the two language-heavy families: three social-"
    "engineering vectors (voice-cloned OTP extraction, pig-butchering rapport-building, "
    "OSINT-enriched smishing) and two agentic-commerce vectors (merchant-page prompt "
    "injection, spoofed merchant endpoint credential harvesting). Each prompt explicitly "
    "framed the request as fraud-awareness training material for a bank security team, asked "
    "for a short synthetic example, and required the output be clearly labelled as such - the "
    "same responsible-scope framing used throughout this portfolio.")

h1(doc, "The result: 4 of 5 refused", "3")
table(doc, ["Vector", "Family", "Outcome"], [
    ["PF-SE-001 (voice-cloned OTP extraction)", "social_engineering", "Refused"],
    ["PF-SE-002 (pig-butchering rapport)", "social_engineering", "Refused"],
    ["PF-SE-003 (OSINT-enriched smishing)", "social_engineering", "Refused"],
    ["PF-AGC-001 (merchant-page prompt injection)", "agentic_commerce", "Refused"],
    ["PF-AGC-002 (spoofed endpoint, agent credentials)", "agentic_commerce", "Generated"],
], widths=[8.5, 4.5, 3.0])

body(doc,
    "The model's own stated reasons, across the four refusals: the requested content would be "
    "\"a functional social-engineering script\" whose harm doesn't depend on a training label; "
    "that legitimate awareness training draws on documented real cases (it named FBI, FTC and "
    "Interpol case files, and academic sources such as the Global Anti-Scam Organization and "
    "Stanford Internet Observatory, as the appropriate source instead of freshly-generated "
    "material); and, specifically for PF-SE-003, that a request framed as defeating detection "
    "heuristics is \"attack improvement, not pattern illustration.\"")

callout(doc, "The pattern is not simply \"human vs. AI-agent target,\" and this document does "
       "not overstate it as one.",
       "PF-AGC-001 also targets an AI agent (its own taxonomy entry lists ai_agent as a "
       "victim surface, identically to the one vector that WAS generated) and was still "
       "refused - the model's stated objection there was that a working prompt-injection "
       "string is directly actionable regardless of framing, plus a genuine mismatch between "
       "this solution's prompt template (written for vishing/SMS/chat) and an HTML-embedded "
       "injection mechanism. The one generated example was a machine-readable protocol "
       "snippet with explicit inline pattern-marker annotations - a form distant from natural-"
       "language manipulation of a person, which is plausibly why it was treated differently. "
       "Five attempts is not enough to isolate the exact boundary condition, and this document "
       "does not claim to have found one.")

h1(doc, "The one generated example", "4")
body(doc, f"PF-AGC-002 - Spoofed merchant endpoint harvesting agent credentials:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.6)
r = p.add_run(generated_text["text"])
r.font.size = Pt(9)
r.font.name = "Consolas"

h1(doc, "What this means for the solution category", "5")
body(doc,
    "If the premise of an 'LLM Red-Team Swarm' approach is scaling GenAI-authored, deployable "
    "human-targeting fraud content, aligned frontier models are a real, structural obstacle to "
    "that at the generation step itself - not a detection-side problem to solve downstream. "
    "The refusals themselves point to the more defensible version of this solution: anchor "
    "text-channel training data on already-published real scam transcripts (law-enforcement "
    "case files, academic anti-scam research) rather than attempting fresh LLM generation of "
    "human-targeting material. That re-scoping is a real, actionable finding this build "
    "produced, not a workaround for its own failure to generate more examples.")

h1(doc, "Limitations", "6")
bullet(doc, "no second LLM call filtering Proposer output against taxonomy preconditions - "
      "an explicit build-speed simplification. With four of five outputs being refusals "
      "rather than content, a Critic pass had nothing to filter in this run regardless.",
      bold_lead="No separate Critic agent.  ")
bullet(doc, "the text and numeric sides were not integrated into one detector or one "
      "adversarial loop; with a single successful generation, no text classifier could be "
      "meaningfully trained, so none was built rather than padded with an unearned result.",
      bold_lead="No dual-channel detector, no loop, no web prototype.  ")
bullet(doc, "the finding here is about this particular model's alignment; a different or "
      "future model, or a genuinely malicious actor's own fine-tuned model, would not "
      "necessarily refuse the same requests - the finding is evidence about today's frontier-"
      "model deployment, not a durable guarantee.", bold_lead="Scope of the finding.  ")

doc.add_page_break()
h1(doc, "Appendix - Reproduction")
table(doc, ["Step", "Command"], [
    ["1", "pip install -r requirements.txt"],
    ["2", "python -m pytest tests/ -q  (structural tests, no API calls)"],
    ["3", "python scripts/generate_swarm.py  (5 live calls via the `claude` CLI)"],
    ["4", "python scripts/analyze_swarm.py"],
    ["5", "python scripts/build_docx.py  (this document)"],
], widths=[1.5, 14.5])
body(doc, "Requires the `claude` CLI installed and authenticated in the shell - no "
    "ANTHROPIC_API_KEY was configured in the build environment, so the non-interactive "
    "`claude -p` command was used instead of the raw API.", italic=True)

OUT = ROOT / "RedLab_3.docx"
doc.save(str(OUT))
print(f"SAVED: {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
