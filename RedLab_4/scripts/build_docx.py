"""Build RedLab_4's full solution walkthrough - detailed, matching RedLab_1's
depth, per explicit request rather than this portfolio's usual lean scope."""
import json
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
CHARTS = ROOT / "artifacts" / "docx_assets"
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2F, 0x5D, 0xC7)
GRAY = RGBColor(0x5B, 0x66, 0x77)
RED = RGBColor(0xB0, 0x2F, 0x2A)
GREEN = RGBColor(0x21, 0x7A, 0x46)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def h1(doc, text, number=None):
    p = doc.add_heading(level=1)
    r = p.add_run((f"{number}. " if number else "") + text)
    r.font.color.rgb = NAVY
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    p.add_run(text).font.color.rgb = ACCENT
    return p


def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def callout(doc, label, text, color=ACCENT, fill="EFF3FA"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(text)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p._p.get_or_add_pPr().append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, "1B2A4A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            if cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def picture(doc, filename, width_cm=14, caption=None):
    path = CHARTS / filename
    if not path.exists():
        body(doc, f"[chart missing: {filename}]", italic=True, color=GRAY)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY


def page_break(doc):
    doc.add_page_break()


# --------------------------------------------------------------------------
from redlab.taxonomy.loader import Taxonomy

tax = Taxonomy.load()
fid_naive = json.loads((ROOT / "artifacts" / "fidelity_naive.json").read_text())
fid_vae = json.loads((ROOT / "artifacts" / "fidelity_vae.json").read_text())
ae_eval = json.loads((ROOT / "artifacts" / "defend_ae_eval.json").read_text())["autoencoder"]
gbm_eval = json.loads((ROOT / "artifacts" / "defend_gbm_eval.json").read_text())["supervised"]
loop_eval = json.loads((ROOT / "artifacts" / "loop_eval.json").read_text())

d_naive = next(m["value"] for m in fid_naive["metrics"] if m["name"] == "discriminator_auc")
d_vae = next(m["value"] for m in fid_vae["metrics"] if m["name"] == "discriminator_auc")
n_projected = ae_eval and json.loads((ROOT/"artifacts"/"defend_ae_eval.json").read_text())["n_projected_vectors"]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = sec.bottom_margin = Cm(2.0)

# ---- Title page ----
for _ in range(4):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RedLab_4")
r.font.size = Pt(40)
r.bold = True
r.font.color.rgb = NAVY

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Adversarial Generative Modeling")
r.font.size = Pt(18)
r.font.color.rgb = ACCENT

st2 = doc.add_paragraph()
st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run("Solution Walkthrough")
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = GRAY

for _ in range(2):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Mastercard Innovation Challenge @ GFF 2026\n"
                "One entry in a multi-solution portfolio - see RedLab_1 (primary), "
                "RedLab_2, RedLab_3, RedLab_6 and RedLab_7 for the other entries")
r.font.size = Pt(11)
r.font.color.rgb = GRAY

for _ in range(4):
    doc.add_paragraph()
box = doc.add_paragraph()
box.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = box.add_run(
    f"discriminator AUC: naive {d_naive:.3f} vs. trained VAE {d_vae:.3f}  ·  "
    f"unsupervised recall {ae_eval['recall_at_0.5fpr']*100:.0f}% vs. supervised "
    f"{gbm_eval['recall_at_0.5fpr']*100:.0f}%  ·  white-box evasion up to "
    f"{max(r['evasion_after'] for r in loop_eval)*100:.0f}%")
r.font.size = Pt(10.5)
r.font.color.rgb = ACCENT
r.italic = True
page_break(doc)

# ---- Contents ----
h1(doc, "Contents")
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fld_begin = OxmlElement("w:fldChar")
fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fld_sep = OxmlElement("w:fldChar")
fld_sep.set(qn("w:fldCharType"), "separate")
fld_text = OxmlElement("w:t")
fld_text.text = "Right-click and choose 'Update Field' to populate the table of contents."
fld_end = OxmlElement("w:fldChar")
fld_end.set(qn("w:fldCharType"), "end")
r_element = run._r
r_element.append(fld_begin)
r_element.append(instr)
r_element.append(fld_sep)
t_run = OxmlElement("w:r")
t_run.append(fld_text)
r_element.addnext(t_run)
t_run.addnext(fld_end)
page_break(doc)

# ==========================================================================
h1(doc, "Executive Summary")
body(doc,
    "This solution tests the assumption most teams entering a challenge like this would "
    "reach for by default: that a deep generative model (a GAN, a VAE, a diffusion model) is "
    "the obvious way to synthesise realistic fraud. Rather than assert that, it builds the "
    "fidelity-measurement harness FIRST, proves the harness actually catches bad synthesis "
    "against a naive baseline, then trains a real conditional VAE and reports its honest "
    "number against that same harness - whatever that number turns out to be.")
body(doc,
    "It pairs that generator with the detection philosophy the brief's own premise calls "
    "for: if the challenge is about NOVEL, unseen fraud, a purely unsupervised detector - "
    "one that never sees a fraud label during training - is a more natural fit than a "
    "supervised classifier. A supervised gradient-boosted baseline is trained on the "
    "identical held-out mechanism purely to quantify what labels buy over none.")
body(doc,
    "Finally, it closes the loop with a genuinely different mechanism from every other "
    "solution in this portfolio: a TRUE white-box, gradient-based adversarial attack, where "
    "the generator's loss function incorporates the frozen detector's own reconstruction-"
    "error computation directly. This is a different threat model from RedLab_1's and "
    "RedLab_7's black-box, query-only attacker - not a replacement for it, a complement.")

table(doc, ["Metric", "Value", "What it means"], [
    ["Naive baseline discriminator AUC", f"{d_naive:.4f}",
     "Independent-marginal sampling; proves the harness catches bad synthesis"],
    ["Trained VAE discriminator AUC", f"{d_vae:.4f}",
     "0.5 = indistinguishable from reference; reported honestly, not tuned to flatter"],
    ["Unsupervised autoencoder recall @ 0.5% FPR", f"{ae_eval['recall_at_0.5fpr']*100:.1f}%",
     "Zero fraud-label exposure, on a held-out never-seen mechanism"],
    ["Supervised GBM recall @ 0.5% FPR", f"{gbm_eval['recall_at_0.5fpr']*100:.1f}%",
     "Same held-out mechanism, trained with labels"],
    ["White-box evasion, before -> after gradient fine-tuning",
     f"{sum(r['evasion_before'] for r in loop_eval)/5*100:.0f}% -> "
     f"{sum(r['evasion_after'] for r in loop_eval)/5*100:.0f}%",
     "Mean across the five synthetic-identity vectors, with bounded drift"],
], widths=[6.0, 4.0, 6.0])

page_break(doc)
print("title + exec summary written")

# ==========================================================================
h1(doc, "Identify — Weighted Toward Synthetic Identity", "1")
body(doc,
    "The taxonomy (42 vectors, unchanged from RedLab_1) is reused wholesale, with no schema "
    "extension - this solution's differentiation is entirely architectural, in Generate and "
    "Defend. The five synthetic-identity vectors were selected as the conditioning targets "
    "because this is the family most naturally about generating a convincing FAKE ARTIFACT - "
    "a document, a face, a persona's spending history - rather than a behavioural sequence, "
    "which is exactly what a tabular generative model is suited to.")
table(doc, ["Vector", "Mechanism", "Maturity"], [
    ["PF-SID-001", "Generated KYC document pack, internally consistent across the set",
     "emerging"],
    ["PF-SID-002", "Camera-injection face morph defeating liveness detection", "emerging"],
    ["PF-SID-003", "Agent-nurtured synthetic identity, months of patient bust-out staging",
     "projected"],
    ["PF-SID-004", "Deepfake video-KYC session, real-time synthesis", "emerging"],
    ["PF-SID-005", "Persona-managed mule recruitment at industrial scale", "emerging"],
], widths=[2.8, 10.2, 3.0])
body(doc, "One of these five, PF-SID-003, carries 'projected' maturity - never observed in "
    "the wild, capability-only. It is one of the vectors held out entirely from both "
    "detectors' training in Section 3, making it the sharpest single test of whether either "
    "detection approach generalises to a genuinely unseen mechanism.")

page_break(doc)

# ==========================================================================
h1(doc, "Generate — Fidelity Harness First, Then a Trained Model", "2")
body(doc,
    "The brief's fidelity criterion is easy to overclaim. The discipline here, taken "
    "directly from the brief's own instruction for this solution: build the measurement "
    "harness BEFORE trusting any generator, prove it catches bad synthesis on a naive "
    "baseline, and only then train and honestly score the real model.")

h2(doc, "2.1 Data anchoring")
body(doc,
    "Anchored on IBM's TabFormer credit-card corpus (24.4M rows), the same reference "
    "RedLab_1 established and this portfolio has reused since - obtained via the public "
    "Git-LFS endpoint without credentials, SHA-256 verified. IBM documents this corpus as "
    "itself synthetic; it is used here as a non-circular external reference for "
    "distributional comparison, not as ground truth. Every number below states plainly what "
    "it was measured against.")

h2(doc, "2.2 The naive baseline, and why it exists")
body(doc,
    "Before training anything, a generator that samples amount, hour, category and channel "
    "INDEPENDENTLY from their correct real marginals - matching each one exactly, but with "
    "zero joint structure - was scored against the harness. This is not a strawman: it is "
    "the specific failure mode RedLab_1's own history recorded a marginals-matched generator "
    "falling into (discriminator AUC 0.847 there), and confirming the harness still catches "
    f"it here (AUC {d_naive:.3f}) is what makes every subsequent number trustworthy.")

h2(doc, "2.3 The trained conditional VAE")
body(doc,
    "A compact conditional VAE (two hidden layers, 64/32 units, a 12-dimensional latent "
    "space - deliberately small, discussed in Section 7) trained on the full combined "
    "corpus, conditioned on attack-vector identity via a learned embedding (43 classes: "
    "42 taxonomy vectors plus a dedicated LEGIT token), so one trained model can sample any "
    "vector, or the legitimate population, on demand. Training converged in 99 seconds over "
    "15 epochs on the full 725,757-row corpus.")

picture(doc, "fidelity_comparison.png", width_cm=12.5,
       caption="Discriminator AUC: naive baseline vs. the trained VAE, both measured "
              "against the identical reference-corpus discriminator test.")
callout(doc, "The honest result: the trained VAE did not clearly beat the naive baseline.",
       f"Naive: {d_naive:.3f}. Trained VAE: {d_vae:.3f}. Both are firmly in the 'weak - "
       f"clear systematic differences' band, not the 'indistinguishable' one. This is "
       f"reported exactly as measured, per the brief's own instruction for this solution: "
       f"do not tune the harness to flatter the model. A small VAE trained for 15 epochs on "
       f"a 4-feature representation (amount, hour, category, channel) captures real "
       f"conditional structure - the loss curves in Section 2.4 show genuine convergence - "
       f"but that does not automatically translate into distributional fidelity superior to "
       f"correctly-sampled independent marginals on THIS metric. The more likely explanation "
       f"is architecture scale and feature scope, not a fundamental flaw in the approach - "
       f"see Section 7's limitations for what a larger, richer-featured version would need.")

picture(doc, "fidelity_per_metric.png", width_cm=13,
       caption="Per-metric comparison: the VAE's advantage, where it has one, is concentrated "
              "in specific marginals rather than uniform across all of them.")

h2(doc, "2.4 Training diagnostics")
body(doc,
    "Reconstruction loss on amount/hour (recon_numeric) fell from 0.120 to 0.014 over 15 "
    "epochs; category reconstruction (recon_mcc) fell from 0.485 to 0.003 - the model learns "
    "to reconstruct conditioning-appropriate categories almost perfectly, which is expected "
    "given category is close to deterministic conditional on attack-vector identity in this "
    "taxonomy. The KL term stabilised around 1.06, indicating the latent space did not "
    "collapse to the prior (a common VAE failure mode) nor diverge unstably.")

page_break(doc)

# ==========================================================================
h1(doc, "Defend — Zero-Label-Exposure Autoencoder vs. Supervised GBM", "3")
body(doc,
    "The primary detector is unsupervised by design: a small autoencoder (48/16 hidden "
    "units, 8-dimensional latent space) trained ONLY on legitimate transactions, using "
    "reconstruction error as the anomaly score, with zero fraud-label exposure at any point "
    "in training. A supervised LightGBM baseline on the identical causal features RedLab_1 "
    "established is trained purely as a comparison point.")
body(doc,
    f"Both are evaluated on an identical held-out slice: {n_projected} of 42 vectors rated "
    f"'projected' maturity are removed from training entirely, at the row level, for BOTH "
    f"detectors - the same zero-exposure discipline RedLab_6 established for exactly this "
    f"purpose, reused here rather than re-derived.")

picture(doc, "defend_comparison.png", width_cm=13,
       caption="Detection efficacy on never-trained-on mechanisms: unsupervised autoencoder "
              "vs. supervised GBM, identical holdout.")
table(doc, ["Metric", "Autoencoder (0 labels)", "Supervised GBM"], [
    ["ROC-AUC", f"{ae_eval['roc_auc']:.4f}", f"{gbm_eval['roc_auc']:.4f}"],
    ["PR-AUC", f"{ae_eval['pr_auc']:.4f}", f"{gbm_eval['pr_auc']:.4f}"],
    ["Recall @ 0.5% FPR", f"{ae_eval['recall_at_0.5fpr']*100:.1f}%",
     f"{gbm_eval['recall_at_0.5fpr']*100:.1f}%"],
    ["Precision @ 0.5% FPR", f"{ae_eval['precision']*100:.1f}%",
     f"{gbm_eval['precision']*100:.1f}%"],
], widths=[6.0, 5.0, 5.0])

callout(doc, "A finding that replicates across two different unsupervised architectures.",
       f"RedLab_6 measured an Isolation-Forest-plus-PCA ensemble reaching 4.0% recall "
       f"against a supervised baseline's 98.2%, on the same style of unseen-mechanism "
       f"holdout. This solution's autoencoder - a genuinely different architecture, deep "
       f"rather than classical - reaches {ae_eval['recall_at_0.5fpr']*100:.1f}% against the "
       f"supervised baseline's {gbm_eval['recall_at_0.5fpr']*100:.1f}%. Two independent "
       f"unsupervised approaches, built separately, land in the same regime: supervised "
       f"models generalise to unseen fraud mechanisms far better than purely legitimate-"
       f"shaped anomaly detection does, in this taxonomy. That consistency across "
       f"architectures is stronger evidence than either result alone.", color=RED,
       fill="FBEDEC")

page_break(doc)

# ==========================================================================
h1(doc, "The Closed Loop — White-Box Gradient-Based Evasion", "4")
body(doc,
    "This is the piece that makes the closed loop literal rather than metaphorical. The "
    "generator's decoder and the detector's autoencoder were deliberately built to share "
    "one continuous representation - numeric features plus category and channel as softmax "
    "probability vectors, never a discrete embedding lookup on either side. That shared "
    "space is what makes a genuine, differentiable generator-to-detector pipeline possible: "
    "the generator's own soft outputs feed directly into the frozen detector's encoder, with "
    "no non-differentiable argmax breaking the gradient anywhere in between.")

h2(doc, "4.1 The threat model, stated explicitly")
body(doc,
    "This is a WHITE-BOX attacker: gradient access to a differentiable generator, and to a "
    "frozen copy of the deployed detector's exact weights. That is a strictly stronger "
    "assumption than RedLab_1's and RedLab_7's BLACK-BOX, query-only attacker (decline-"
    "response oracle only, no internals). Neither threat model is more realistic than the "
    "other - a real portfolio benefits from having both, and this document does not claim "
    "otherwise. The evasion numbers below are not comparable to RedLab_1's or RedLab_7's: "
    "different attacker, different defender, different question.")

h2(doc, "4.2 Mechanism")
body(doc,
    "For each of the five synthetic-identity vectors, 500 conditioned samples are drawn "
    "from the generator's latent space. A free latent vector per sample is then optimised "
    "for 200 gradient steps (Adam, lr=0.02) against a loss combining the frozen detector's "
    "reconstruction error (minimise - be judged 'normal') with an L2 drift penalty against "
    "the sample's own un-optimised decode (stay close to the originally-conditioned, "
    "realistic output). This is the continuous-optimisation analogue of RedLab_1's "
    "evasion_rate x value_retention fitness function: evasion is only counted as a win if "
    "the sample has not drifted arbitrarily far from what the generator was actually asked "
    "to produce.")

picture(doc, "loop_evasion.png", width_cm=13,
       caption="Evasion rate before and after white-box gradient fine-tuning, per "
              "synthetic-identity vector.")
table(doc, ["Vector", "Evasion before", "Evasion after", "Mean drift"], [
    [r["vector_id"], f"{r['evasion_before']*100:.1f}%", f"{r['evasion_after']*100:.1f}%",
     f"{r['mean_drift']:.3f}"] for r in loop_eval
], widths=[3.0, 4.0, 4.0, 4.0])

callout(doc, "Read this result alongside Section 3, not in isolation.",
       f"Baseline evasion is already high (~90%) before any adversarial optimisation, "
       f"because the autoencoder detector's overall recall on unseen mechanisms is "
       f"already weak ({ae_eval['recall_at_0.5fpr']*100:.1f}%, from Section 3) - most "
       f"unoptimised synthetic-identity samples were already evading it. White-box "
       f"fine-tuning closes most of what remains: the detector's catch rate on these five "
       f"vectors drops from roughly 10% to roughly 1%, with mean drift held under 0.05 in "
       f"the shared feature space, so the evading samples have not been optimised into "
       f"something unrecognisable as the conditioned attack vector. The loop's real "
       f"contribution here is closing the LAST gap, not creating the vulnerability - the "
       f"vulnerability is already present in Section 3's weak unsupervised recall.")

page_break(doc)

# ==========================================================================
h1(doc, "Working Prototype", "5")
body(doc, "Not built for this solution - an explicit scope decision. The most valuable "
    "screen this solution would justify is a live view of the evasion optimisation itself: "
    "the detector's reconstruction-error distribution for a batch of conditioned samples, "
    "animated across gradient steps, showing the mass of samples crossing the threshold in "
    "real time. That is a natural extension, not attempted here in the interest of build "
    "time relative to the rest of this portfolio.")

page_break(doc)

# ==========================================================================
h1(doc, "Real-World Feasibility in Live Payment Environments", "6")
bullet(doc, "the shared continuous representation between generator and detector is not "
      "free - a production system running this architecture needs its generator and "
      "detector co-versioned and co-deployed, since the adversarial loop's differentiable "
      "path depends on both operating in identical feature space. Redeploying one without "
      "the other silently breaks the loop, with no error at inference time.",
      bold_lead="Architectural coupling has an operational cost.  ")
bullet(doc, "Section 3's finding is the central feasibility point: an unsupervised detector "
      "alone is not adequate for unseen-mechanism fraud in this taxonomy, replicated across "
      "two independent architectures now (RedLab_6's classical ensemble, this solution's "
      "deep autoencoder). A production system should treat unsupervised anomaly scoring as "
      "a secondary, complementary signal - exactly the framing RedLab_1's C4 'novelty "
      "channel' used - not a standalone answer to the challenge's 'detects novel fraud' "
      "premise.", bold_lead="Unsupervised-only is not production-ready here.  ")
bullet(doc, "a white-box attacker with weight access to a deployed model is a real, if less "
      "commonly discussed, threat: model weights leak, get extracted via distillation, or "
      "are exposed through an insecurely-deployed inference endpoint. Section 4's finding - "
      "that gradient access closes an evasion gap almost completely, with minimal drift from "
      "a realistic sample - is a concrete argument for weight confidentiality and endpoint "
      "hardening as a fraud-control measure, not just a conventional security one.",
      bold_lead="Weight confidentiality is a fraud control, not just a security one.  ")

page_break(doc)

# ==========================================================================
h1(doc, "Responsible Red-Teaming and Limitations", "7")

h2(doc, "7.1 Scope boundary")
body(doc, "Identical to the rest of this portfolio: everything synthetic and sandboxed, no "
    "real personal data, no real account or card numbers, no deployable tooling against any "
    "live rail.")

h2(doc, "7.2 Limitations")
bullet(doc, "the generator operates on four features (amount, hour, category, channel) - "
      "the same four the fidelity harness scores. RedLab_1's causal feature set has 30. A "
      "generator conditioned on a richer feature space, with a larger latent dimension and "
      "more training capacity, is the most likely lever to close the gap to the naive "
      "baseline that Section 2.3 reports honestly rather than papering over.",
      bold_lead="Small model, narrow feature scope.  ")
bullet(doc, "GAN and full diffusion architectures were considered and explicitly not "
      "attempted, given this portfolio's own prior experience with unstable, long-running "
      "training loops (documented in RedLab_2's build history) and the smoke-test discipline "
      "adopted since. A VAE was chosen specifically because it trains stably to convergence "
      "in under two minutes on this hardware, verified before the real run was ever "
      "launched.", bold_lead="VAE, not GAN or diffusion.  ")
bullet(doc, "the white-box loop optimises a free latent vector per sample rather than "
      "fine-tuning the generator's own weights end-to-end - a lighter-weight approximation "
      "that is still a genuine gradient-based attack against the frozen detector, but is not "
      "identical to updating the generator's parameters directly. Both are legitimate "
      "white-box constructions; this document does not claim to have implemented the "
      "heavier variant.", bold_lead="Latent optimisation, not full generator fine-tuning.  ")
bullet(doc, "several long-running training invocations during this build entered an "
      "unrecoverable OS-level process state (documented in this portfolio's own git "
      "history for RedLab_2 and encountered again here) when run through certain shell "
      "invocation patterns. The final, reliable pattern used throughout this solution - "
      "`nohup ... & disown`, with results polled via output-file existence rather than "
      "process exit - is itself a finding worth carrying into any future solution in this "
      "portfolio that trains a neural model.", bold_lead="A build-environment finding, stated plainly.  ")

page_break(doc)

# ==========================================================================
h1(doc, "Appendix — Reproduction Guide")
table(doc, ["Step", "Command", "Produces"], [
    ["1", "pip install -r requirements.txt", "Dependencies"],
    ["2", "python -m pytest tests/ -q", "10 tests validating the shared representation, "
     "fidelity harness sanity, and loop invariants"],
    ["3", "python scripts/train_cvae.py", "The trained conditional VAE (~100s)"],
    ["4", "python scripts/save_naive_fidelity.py\npython scripts/eval_fidelity.py",
     "Naive baseline and VAE fidelity reports"],
    ["5", "python scripts/train_ae_only.py\npython scripts/train_gbm_only.py",
     "The unsupervised and supervised detectors"],
    ["6", "python scripts/run_loop.py", "The white-box evasion results"],
    ["7", "python scripts/build_charts.py\npython scripts/build_docx.py",
     "This document, regenerated from the artifacts above"],
], widths=[1.2, 8.3, 6.5])
body(doc, "All long-running scripts in this reproduction guide should be launched with "
    "`nohup python3 script.py > out.log 2>&1 & disown` and their completion checked via the "
    "output artifact's existence, per Section 7.2's build-environment finding.", italic=True)

OUT_PATH = ROOT / "RedLab_4.docx"
doc.save(str(OUT_PATH))
print(f"\nSAVED: {OUT_PATH}  ({OUT_PATH.stat().st_size/1024:.0f} KB, "
     f"{len(doc.paragraphs)} paragraphs)")
