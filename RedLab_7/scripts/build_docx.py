"""Build RedLab_7's short-form solution walkthrough. Deliberately compact,
matching this solution's scope."""
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2F, 0x5D, 0xC7)
GRAY = RGBColor(0x5B, 0x66, 0x77)
RED = RGBColor(0xB0, 0x2F, 0x2A)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def h1(doc, text, number=None):
    p = doc.add_heading(level=1)
    r = p.add_run((f"{number}. " if number else "") + text)
    r.font.color.rgb = NAVY
    return p


def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def callout(doc, label, text, color=ACCENT, fill="EFF3FA"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(text)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p._p.get_or_add_pPr().append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, "1B2A4A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            if cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


# --------------------------------------------------------------------------
from redlab.taxonomy.loader import Taxonomy

tax = Taxonomy.load()
ev = json.loads((ROOT / "artifacts" / "eval_results.json").read_text())
spread = ev["spread"]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = sec.bottom_margin = Cm(2.0)

for _ in range(5):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RedLab_7")
r.font.size = Pt(38)
r.bold = True
r.font.color.rgb = NAVY
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Federated Consortium Defense")
r.font.size = Pt(17)
r.font.color.rgb = ACCENT
st2 = doc.add_paragraph()
st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run("Solution Walkthrough (short form)")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = GRAY
for _ in range(2):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Mastercard Innovation Challenge @ GFF 2026\n"
                "One entry in a multi-solution portfolio - see RedLab_1 (primary), "
                "RedLab_2 and RedLab_6 for the other entries")
r.font.size = Pt(10.5)
r.font.color.rgb = GRAY
doc.add_page_break()

# ==========================================================================
h1(doc, "Summary")
body(doc,
    "RedLab_1's own taxonomy already states, in the detection hypotheses for two of its "
    "vectors, that mule-account layering and distributed BIN testing work BECAUSE any single "
    "issuer or acquirer sees only a fragment of the attack - the stated mitigation for one is "
    "literally 'network-level (not merchant-level) monitoring.' This solution tests that claim "
    "directly: partition the population across four synthetic institutions, and measure how "
    "much detection capability a privacy-preserving federated model recovers versus an "
    "unrealistic fully-centralized one, against the realistic status quo of institutions that "
    "never collaborate at all.")
body(doc,
    "Scoped deliberately small, same discipline as RedLab_6: Identify and Generate are reused "
    "wholesale from RedLab_1. The new work is entirely in Defend - institution partitioning, "
    "local-only models, one round of federated coefficient averaging, and a centralized-oracle "
    "upper bound, compared on the identical cross-institution test set. No adversarial loop, "
    "no web prototype - stated here as a scope decision, not an incomplete result.")

table(doc, ["Metric", "Value"], [
    ["Synthetic institutions", "4 (issuer-side and acquirer-side, independently assigned)"],
    ["Cross-institution hero vectors",
     "PF-IND-005 (mule layering), PF-CT-001 (distributed BIN testing)"],
    ["PF-IND-005 spread", f"{spread['PF-IND-005']['n_issuers']} of 4 issuers, "
     f"{spread['PF-IND-005']['n_acquirers']} of 4 acquirers"],
    ["PF-CT-001 spread", f"{spread['PF-CT-001']['n_issuers']} of 4 issuers, "
     f"{spread['PF-CT-001']['n_acquirers']} of 4 acquirers"],
    ["Local-only recall @ 1% FPR", f"{ev['local']['recall_at_1fpr']*100:.1f}%"],
    ["Federated recall @ 1% FPR", f"{ev['federated']['recall_at_1fpr']*100:.1f}%"],
    ["Centralized-oracle recall @ 1% FPR", f"{ev['centralized_oracle']['recall_at_1fpr']*100:.1f}%"],
    ["Federation recovers", f"{ev['recovery_pct']:.0f}% of the local-to-oracle gap"],
], widths=[7.5, 8.5])

h1(doc, "What's reused vs. what's new", "1")
bullet(doc, "taxonomy (42 vectors, unchanged), calibrated world simulator, attack injection, "
      "and the 30 causal features - copied from RedLab_1 without modification.",
      bold_lead="Reused.  ")
bullet(doc, "institution partitioning (redlab/sim/institutions.py), the federation mechanics "
      "(redlab/defend/federated.py), and the local/federated/oracle comparison.",
      bold_lead="New.  ")

h1(doc, "Institution partitioning", "2")
body(doc,
    "Every cardholder is assigned an issuer, every merchant an acquirer, via independent "
    "deterministic hashes - no change to the world simulator itself, since fragmentation "
    "across institutions is a property of who owns which entity, not of how transactions are "
    "generated. Both hero vectors were verified to genuinely span all four institutions on "
    "both sides once partitioned, confirming the cross-institution blind spot the taxonomy "
    "describes is real in this simulated population, not just asserted.")

h1(doc, "Federation mechanics", "3")
body(doc,
    "A linear (logistic regression) model, not a gradient-boosted one, is used deliberately: "
    "averaging coefficient vectors across institutions is FedAvg's actual, well-defined "
    "mechanic. There is no equivalent standard operation for averaging two ensembles of trees. "
    "Each institution trains locally to convergence; coefficients are averaged once, weighted "
    "by each institution's sample count - the simplest case of FedAvg (one communication "
    "round), stated as a simplification rather than the full multi-round protocol a production "
    "system would run. Only coefficients and an intercept cross the institutional boundary; "
    "the federated_average function's signature structurally cannot accept a raw training "
    "frame, which is the actual privacy property this solution claims - and is checked "
    "directly by an automated test, not just asserted in a docstring.")

h1(doc, "The headline result, and why the gap is small", "4")
table(doc, ["Model", "ROC-AUC", "PR-AUC", "Recall @ 1% FPR"], [
    ["Local-only (status quo)", f"{ev['local']['roc_auc']:.4f}", f"{ev['local']['pr_auc']:.4f}",
     f"{ev['local']['recall_at_1fpr']*100:.1f}%"],
    ["Federated", f"{ev['federated']['roc_auc']:.4f}", f"{ev['federated']['pr_auc']:.4f}",
     f"{ev['federated']['recall_at_1fpr']*100:.1f}%"],
    ["Centralized oracle", f"{ev['centralized_oracle']['roc_auc']:.4f}",
     f"{ev['centralized_oracle']['pr_auc']:.4f}",
     f"{ev['centralized_oracle']['recall_at_1fpr']*100:.1f}%"],
], widths=[6.5, 3.0, 3.0, 3.5])
callout(doc, "The finding, reported exactly as measured.",
       f"Federation recovers {ev['recovery_pct']:.0f}% of the gap between local-only and "
       f"centralized - but that gap is small to begin with (local-only already reaches "
       f"{ev['local']['recall_at_1fpr']*100:.1f}% recall). The reason is diagnosable, not "
       f"mysterious: every one of the 30 causal features this solution inherited from RedLab_1 "
       f"is scoped to a single user, merchant, or device (u_*, m_*, d_* prefixes, verified by "
       f"direct inspection) - none of them require visibility across institutions to compute. "
       f"A user's own transaction history is already fully available to their own issuer; a "
       f"merchant's own history is already fully available to its own acquirer. Federation can "
       f"only recover value that exists as a feature no single institution could compute "
       f"alone - something like 'how many distinct institutions has this merchant been paid "
       f"from in the last hour,' which this feature set does not contain. The real lesson is "
       f"about feature design, not about whether federation works: a consortium defence "
       f"earns its complexity only when the features it shares are the kind no single member "
       f"could build alone.")

h1(doc, "Feasibility and limitations", "5")
bullet(doc, "single-round coefficient averaging is FedAvg's simplest case; a production "
      "deployment would run multiple rounds and typically differential-privacy-noised "
      "updates, neither implemented here.", bold_lead="Simplification, stated plainly.  ")
bullet(doc, "the diagnosis in Section 4 is directly actionable: a genuinely federation-"
      "dependent detector needs cross-institution-aware features computed via secure "
      "aggregation (e.g. a merchant's institution-fan-in count, shared without any single "
      "institution seeing another's raw transactions) - this is exactly the kind of feature "
      "real card-network fraud consortiums build, and is the natural next step, not "
      "attempted here in the interest of build time.", bold_lead="What would change the result.  ")
bullet(doc, "no adversarial loop, no web prototype - explicit scope decision for a fast build.",
      bold_lead="Out of scope, by design.  ")

doc.add_page_break()
h1(doc, "Appendix - Reproduction")
table(doc, ["Step", "Command"], [
    ["1", "pip install -r requirements.txt"],
    ["2", "python -m pytest tests/ -q"],
    ["3", "python scripts/evaluate.py  (partitions, fits all three models, compares)"],
    ["4", "python scripts/build_docx.py  (this document)"],
], widths=[1.5, 14.5])
body(doc, "World, taxonomy, and features are generated the same way as RedLab_1 - see that "
    "solution's README for the underlying generation commands.", italic=True)

OUT = ROOT / "RedLab_7.docx"
doc.save(str(OUT))
print(f"SAVED: {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
